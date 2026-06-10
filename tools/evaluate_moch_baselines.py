# -*- coding: utf-8 -*-
import argparse
import csv
import math
import os
from multiprocessing import Pool
from collections import defaultdict
from pathlib import Path

import numpy as np
from PIL import Image
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

try:
    import cv2
except Exception:
    cv2 = None


IMAGE_SUFFIXES = (".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff")


def sorted_images(path):
    return sorted(p for p in Path(path).iterdir() if p.suffix.lower() in IMAGE_SUFFIXES)


def list_moch_gt_frames(root, splits):
    rows = []
    root = Path(root)
    for split in splits:
        split_root = root / "data" / split
        if not split_root.exists():
            continue
        for scene_dir in sorted(p for p in split_root.iterdir() if p.is_dir()):
            gt_dir = scene_dir / "gts"
            img_dir = scene_dir / "images"
            if not gt_dir.exists() or not img_dir.exists():
                continue
            gt_files = sorted_images(gt_dir)
            img_files = sorted_images(img_dir)
            img_by_stem = {p.stem: p for p in img_files}
            for idx, gt_path in enumerate(gt_files):
                if gt_path.stem not in img_by_stem:
                    continue
                rows.append(
                    {
                        "dataset": "MoCH",
                        "split": split,
                        "scene": scene_dir.name,
                        "frame_index": idx,
                        "frame_count": len(gt_files),
                        "frame_name": gt_path.stem,
                        "pred_name": gt_path.stem,
                        "gt_path": str(gt_path),
                        "img_path": str(img_by_stem[gt_path.stem]),
                    }
                )
    return rows


def read_gray(path, size=None):
    if cv2 is not None:
        img = cv2.imread(str(path), cv2.IMREAD_GRAYSCALE)
        if img is None:
            raise FileNotFoundError(path)
        if size is not None and (img.shape[1], img.shape[0]) != size:
            img = cv2.resize(img, size, interpolation=cv2.INTER_LINEAR)
        return img.astype(np.float32)
    with Image.open(path) as raw:
        img = raw.convert("L")
        if size is not None and img.size != size:
            img = img.resize(size, Image.BILINEAR)
        return np.asarray(img, dtype=np.float32)


def safe_div(num, den):
    return float(num / den) if den else 0.0


def binary_erosion(mask):
    if cv2 is not None:
        kernel = np.ones((3, 3), dtype=np.uint8)
        return cv2.erode(mask.astype(np.uint8), kernel, iterations=1).astype(bool)
    if mask.shape[0] < 3 or mask.shape[1] < 3:
        return np.zeros_like(mask, dtype=bool)
    out = np.zeros_like(mask, dtype=bool)
    center = mask[1:-1, 1:-1]
    out[1:-1, 1:-1] = (
        center
        & mask[:-2, 1:-1]
        & mask[2:, 1:-1]
        & mask[1:-1, :-2]
        & mask[1:-1, 2:]
        & mask[:-2, :-2]
        & mask[:-2, 2:]
        & mask[2:, :-2]
        & mask[2:, 2:]
    )
    return out


def binary_dilation(mask, iterations=1):
    if cv2 is not None:
        kernel = np.ones((3, 3), dtype=np.uint8)
        return cv2.dilate(mask.astype(np.uint8), kernel, iterations=iterations).astype(bool)
    out = mask.astype(bool)
    for _ in range(iterations):
        padded = np.pad(out, 1, mode="constant", constant_values=False)
        nxt = np.zeros_like(out, dtype=bool)
        for dy in range(3):
            for dx in range(3):
                nxt |= padded[dy : dy + out.shape[0], dx : dx + out.shape[1]]
        out = nxt
    return out


def component_count(mask):
    if cv2 is not None:
        num, _labels = cv2.connectedComponents(mask.astype(np.uint8), 8)
        return max(0, int(num) - 1)
    h, w = mask.shape
    seen = np.zeros_like(mask, dtype=bool)
    count = 0
    for y in range(h):
        xs = np.where(mask[y] & ~seen[y])[0]
        for x0 in xs:
            if seen[y, x0] or not mask[y, x0]:
                continue
            count += 1
            stack = [(y, int(x0))]
            seen[y, x0] = True
            while stack:
                cy, cx = stack.pop()
                for ny in (cy - 1, cy, cy + 1):
                    for nx in (cx - 1, cx, cx + 1):
                        if ny == cy and nx == cx:
                            continue
                        if 0 <= ny < h and 0 <= nx < w and mask[ny, nx] and not seen[ny, nx]:
                            seen[ny, nx] = True
                            stack.append((ny, nx))
    return count


def bbox(mask):
    ys, xs = np.where(mask)
    if len(xs) == 0:
        return None
    return int(xs.min()), int(ys.min()), int(xs.max()) + 1, int(ys.max()) + 1


def bbox_iou(a, b):
    if a is None or b is None:
        return 0.0
    ax1, ay1, ax2, ay2 = a
    bx1, by1, bx2, by2 = b
    ix1, iy1 = max(ax1, bx1), max(ay1, by1)
    ix2, iy2 = min(ax2, bx2), min(ay2, by2)
    inter = max(0, ix2 - ix1) * max(0, iy2 - iy1)
    area_a = max(0, ax2 - ax1) * max(0, ay2 - ay1)
    area_b = max(0, bx2 - bx1) * max(0, by2 - by1)
    return safe_div(inter, area_a + area_b - inter)


def frame_metrics(pred, gt):
    pred = pred.astype(bool)
    gt = gt.astype(bool)
    inter = np.logical_and(pred, gt).sum()
    union = np.logical_or(pred, gt).sum()
    p_area = int(pred.sum())
    g_area = int(gt.sum())
    iou = safe_div(inter, union)
    dice = safe_div(2 * inter, p_area + g_area)
    precision = safe_div(inter, p_area)
    recall = safe_div(inter, g_area)
    pred_edge = np.logical_xor(binary_dilation(pred), binary_erosion(pred))
    gt_edge = np.logical_xor(binary_dilation(gt), binary_erosion(gt))
    pred_band = binary_dilation(pred_edge, 2)
    gt_band = binary_dilation(gt_edge, 2)
    bp = safe_div(np.logical_and(pred_edge, gt_band).sum(), pred_edge.sum())
    br = safe_div(np.logical_and(gt_edge, pred_band).sum(), gt_edge.sum())
    bf1 = safe_div(2 * bp * br, bp + br)
    return {
        "iou": iou,
        "dice": dice,
        "precision": precision,
        "recall": recall,
        "failure": 1.0 if iou < 0.3 else 0.0,
        "boundary_f1": bf1,
        "component_count": component_count(pred),
        "pred_gt_area_ratio": safe_div(p_area, g_area),
        "bbox_iou": bbox_iou(bbox(pred), bbox(gt)),
    }


def mean_rows(rows, keys):
    if not rows:
        return {k: 0.0 for k in keys}
    return {k: float(np.mean([float(r[k]) for r in rows])) for k in keys}


def write_csv(path, rows):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = list(rows[0].keys()) if rows else []
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def evaluate_row_task(task):
    row, model_specs, threshold = task
    frame_rows = []
    missing = []
    gt = read_gray(row["gt_path"]) >= 128
    denom = max(1, row["frame_count"] - 1)
    pos = row["frame_index"] / denom
    phase = "early" if pos < 1 / 3 else "middle" if pos < 2 / 3 else "late"
    for model, pred_root in model_specs:
        pred_path = Path(pred_root) / "MoCH" / row["split"] / row["scene"] / (row["pred_name"] + ".png")
        if not pred_path.exists():
            missing.append({"model": model, **row, "pred_path": str(pred_path)})
            continue
        pred_prob = read_gray(pred_path, size=(gt.shape[1], gt.shape[0])) / 255.0
        pred = pred_prob >= threshold
        met = frame_metrics(pred, gt)
        frame_rows.append(
            {
                "dataset": "MoCH",
                "model": model,
                "split": row["split"],
                "scene": row["scene"],
                "frame_index": row["frame_index"],
                "frame_count": row["frame_count"],
                "frame_name": row["frame_name"],
                "phase": phase,
                "pred_path": str(pred_path),
                "gt_path": row["gt_path"],
                **met,
            }
        )
    return frame_rows, missing


def evaluate(args):
    gt_rows = list_moch_gt_frames(args.moch_root, args.splits)
    model_specs = []
    for model_arg in args.model:
        name, root = model_arg.split("=", 1)
        model_specs.append((name, Path(root)))

    metric_keys = [
        "iou",
        "dice",
        "precision",
        "recall",
        "failure",
        "boundary_f1",
        "component_count",
        "pred_gt_area_ratio",
        "bbox_iou",
    ]
    frame_rows = []
    missing = []
    worker_specs = [(name, str(root)) for name, root in model_specs]
    tasks = ((row, worker_specs, args.threshold) for row in gt_rows)
    workers = max(1, int(args.workers))
    if workers == 1:
        iterator = map(evaluate_row_task, tasks)
        for idx, (fr, ms) in enumerate(iterator, 1):
            frame_rows.extend(fr)
            missing.extend(ms)
            if idx % 250 == 0:
                print(f"evaluated {idx}/{len(gt_rows)} MoCH GT frames", flush=True)
    else:
        with Pool(processes=workers) as pool:
            for idx, (fr, ms) in enumerate(pool.imap(evaluate_row_task, tasks, chunksize=8), 1):
                frame_rows.extend(fr)
                missing.extend(ms)
                if idx % 250 == 0:
                    print(f"evaluated {idx}/{len(gt_rows)} MoCH GT frames", flush=True)

    model_order = {name: i for i, (name, _root) in enumerate(model_specs)}
    split_order = {name: i for i, name in enumerate(args.splits)}
    frame_rows.sort(key=lambda r: (model_order.get(r["model"], 999), split_order.get(r["split"], 999), r["scene"], int(r["frame_index"])))

    overall_rows = []
    split_rows = []
    sequence_rows = []
    phase_rows = []
    first_rows = []
    for model, _root in model_specs:
        mr = [r for r in frame_rows if r["model"] == model]
        overall_rows.append({"dataset": "MoCH", "model": model, "split": "All", "frames": len(mr), **mean_rows(mr, metric_keys)})
        for split in args.splits:
            sr = [r for r in mr if r["split"] == split]
            split_rows.append({"dataset": "MoCH", "model": model, "split": split, "frames": len(sr), **mean_rows(sr, metric_keys)})
        scenes = sorted({(r["split"], r["scene"]) for r in mr})
        for split, scene in scenes:
            seq = [r for r in mr if r["split"] == split and r["scene"] == scene]
            sequence_rows.append({"dataset": "MoCH", "model": model, "split": split, "scene": scene, "frames": len(seq), **mean_rows(seq, metric_keys)})
        for split in ["All"] + list(args.splits):
            base = mr if split == "All" else [r for r in mr if r["split"] == split]
            for phase in ["early", "middle", "late"]:
                pr = [r for r in base if r["phase"] == phase]
                phase_rows.append({"dataset": "MoCH", "model": model, "split": split, "phase": phase, "frames": len(pr), **mean_rows(pr, metric_keys)})
        first = [r for r in mr if int(r["frame_index"]) == 0]
        first_rows.append({"dataset": "MoCH", "model": model, "split": "All", "frames": len(first), **mean_rows(first, metric_keys)})

    out = Path(args.out_dir)
    out.mkdir(parents=True, exist_ok=True)
    write_csv(out / "frame_metrics.csv", frame_rows)
    write_csv(out / "overall_metrics.csv", overall_rows)
    write_csv(out / "split_metrics.csv", split_rows)
    write_csv(out / "sequence_metrics.csv", sequence_rows)
    write_csv(out / "phase_metrics.csv", phase_rows)
    write_csv(out / "first_frame_metrics.csv", first_rows)
    if missing:
        write_csv(out / "missing_predictions.csv", missing)
    build_docx(args.docx, overall_rows, split_rows, sequence_rows, phase_rows, first_rows, missing, args)


def fnum(value, digits=4):
    try:
        return ("{:.%df}" % digits).format(float(value))
    except Exception:
        return ""


def set_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in doc.styles:
            doc.styles[name].font.name = "Malgun Gothic"
            doc.styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def build_docx(path, overall, splits, sequences, phases, firsts, missing, args):
    doc = Document()
    set_font(doc)
    doc.add_heading("Baseline MoCH 평가 결과", level=0)
    doc.add_paragraph(
        "EMIP와 TSP-SAM baseline을 MoCH 데이터셋의 Train, Validation, Test split 전체 GT frame 기준으로 평가했습니다. "
        "평가는 bbox가 아니라 mask 기준이며, 예측 probability map을 threshold {:.2f}에서 이진화했습니다.".format(args.threshold)
    )
    if missing:
        doc.add_paragraph("누락 prediction {}개가 있어 해당 frame은 평균에서 제외했습니다. 상세 목록은 missing_predictions.csv에 저장했습니다.".format(len(missing)))

    doc.add_heading("지표 의미", level=1)
    add_table(
        doc,
        ["지표", "의미"],
        [
            ["IoU", "예측 mask와 GT mask의 교집합 / 합집합입니다. 높을수록 좋습니다."],
            ["Dice", "겹친 영역을 부드럽게 보는 지표입니다. 높을수록 좋습니다."],
            ["Precision", "예측 foreground 중 실제 foreground 비율입니다. 낮으면 배경까지 많이 잡은 것입니다."],
            ["Recall", "실제 foreground 중 모델이 찾아낸 비율입니다. 낮으면 객체를 놓친 것입니다."],
            ["Failure", "IoU 0.3 미만 frame 비율입니다. 낮을수록 안정적입니다."],
            ["Boundary F1", "예측 경계와 GT 경계의 일치도입니다."],
            ["Component", "예측 mask 조각 수입니다. 높으면 마스크가 많이 쪼개진 것입니다."],
            ["Area", "예측 foreground 면적 / GT foreground 면적입니다. 1에 가까울수록 좋습니다."],
        ],
    )

    doc.add_heading("전체 결과", level=1)
    add_table(
        doc,
        ["모델", "frames", "IoU", "Dice", "Precision", "Recall", "Failure", "Boundary F1", "Component", "Area"],
        [
            [
                r["model"],
                r["frames"],
                fnum(r["iou"]),
                fnum(r["dice"]),
                fnum(r["precision"]),
                fnum(r["recall"]),
                fnum(r["failure"]),
                fnum(r["boundary_f1"]),
                fnum(r["component_count"], 2),
                fnum(r["pred_gt_area_ratio"], 2),
            ]
            for r in overall
        ],
    )
    if overall:
        best = max(overall, key=lambda r: float(r["iou"]))
        doc.add_paragraph(
            "전체 MoCH 기준 최고 IoU는 {}의 {}입니다. 다만 MoCH는 사람 중심 camouflage 영상이라 "
            "면적 보정과 boundary 안정성도 함께 보는 것이 좋습니다.".format(best["model"], fnum(best["iou"]))
        )

    doc.add_heading("Split별 결과", level=1)
    add_table(
        doc,
        ["split", "모델", "frames", "IoU", "Dice", "Recall", "Failure", "Boundary F1", "Area"],
        [
            [r["split"], r["model"], r["frames"], fnum(r["iou"]), fnum(r["dice"]), fnum(r["recall"]), fnum(r["failure"]), fnum(r["boundary_f1"]), fnum(r["pred_gt_area_ratio"], 2)]
            for r in splits
        ],
    )

    doc.add_heading("첫 frame 성능", level=1)
    add_table(
        doc,
        ["모델", "frames", "IoU", "Dice", "Recall", "Failure", "Boundary F1", "Area"],
        [[r["model"], r["frames"], fnum(r["iou"]), fnum(r["dice"]), fnum(r["recall"]), fnum(r["failure"]), fnum(r["boundary_f1"]), fnum(r["pred_gt_area_ratio"], 2)] for r in firsts],
    )

    doc.add_heading("초/중/후반 시간대", level=1)
    add_table(
        doc,
        ["split", "phase", "모델", "frames", "IoU", "Dice", "Failure", "Boundary F1", "Area"],
        [
            [r["split"], r["phase"], r["model"], r["frames"], fnum(r["iou"]), fnum(r["dice"]), fnum(r["failure"]), fnum(r["boundary_f1"]), fnum(r["pred_gt_area_ratio"], 2)]
            for r in phases
            if r["split"] == "All"
        ],
    )

    doc.add_heading("취약 sequence", level=1)
    weak_rows = []
    for model in sorted({r["model"] for r in sequences}):
        rows = [r for r in sequences if r["model"] == model and int(r["frames"]) >= 5]
        rows.sort(key=lambda r: (float(r["iou"]), -float(r["failure"])))
        for r in rows[:10]:
            weak_rows.append([model, r["split"], r["scene"], r["frames"], fnum(r["iou"]), fnum(r["failure"]), fnum(r["boundary_f1"]), fnum(r["pred_gt_area_ratio"], 2)])
    add_table(doc, ["모델", "split", "sequence", "frames", "IoU", "Failure", "Boundary F1", "Area"], weak_rows)

    doc.add_heading("간단 해석", level=1)
    by_model = {r["model"]: r for r in overall}
    if "EMIP" in by_model and "TSP-SAM" in by_model:
        e, t = by_model["EMIP"], by_model["TSP-SAM"]
        better = "EMIP" if float(e["iou"]) >= float(t["iou"]) else "TSP-SAM"
        doc.add_paragraph(
            "MoCH 전체 기준으로는 {}가 IoU 기준 더 높았습니다. EMIP와 TSP-SAM의 차이는 IoU {}, Failure {}입니다.".format(
                better,
                fnum(abs(float(e["iou"]) - float(t["iou"]))),
                fnum(abs(float(e["failure"]) - float(t["failure"]))),
            )
        )
        doc.add_paragraph(
            "Failure가 높거나 Area가 1에서 멀어지는 split/sequence는 모델이 객체를 놓치거나 배경까지 넓게 잡는 구간입니다. "
            "Boundary F1이 낮은 구간은 마스크가 대략 겹치더라도 실제 객체 경계가 흐려진 경우로 해석하면 됩니다."
        )
    doc.save(path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--moch-root", default="MoCH")
    parser.add_argument("--splits", nargs="+", default=["Train", "Validation", "Test"])
    parser.add_argument("--out-dir", default="analysis_results_baseline_moch")
    parser.add_argument("--docx", default="Baseline_MoCH.docx")
    parser.add_argument("--threshold", type=float, default=0.5)
    parser.add_argument("--model", action="append", required=True)
    parser.add_argument("--workers", type=int, default=min(8, os.cpu_count() or 1))
    args = parser.parse_args()
    evaluate(args)


if __name__ == "__main__":
    main()
